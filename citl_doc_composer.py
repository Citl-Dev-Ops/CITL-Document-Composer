#!/usr/bin/env python3
"""
CITL Document Composer v2.0
Professional document authoring tool with built-in screen capture,
PDF export, Apothecary font library, and RTC navy/gold design theme.
"""
from __future__ import annotations

import ctypes
import io
import os
import queue
import re
import shutil
import subprocess
import sys
import threading
import traceback
import winreg
from datetime import date
from pathlib import Path
from typing import List, Optional

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox, scrolledtext, ttk
except ImportError:
    print("tkinter required.")
    sys.exit(1)

from citl_doc_theme import (
    add_body,
    add_callout,
    add_cover_page,
    add_h1_with_bar,
    add_h3,
    add_rule,
    add_screenshot_placeholder,
    apply_citl_styles,
    get_active_doc_style,
    get_doc_style_fonts,
    get_doc_style_names,
    get_missing_fonts,
    set_doc_style,
)
from citl_doc_templates import (
    TEMPLATE_NAMES,
    get_best_model,
    get_best_vision_model,
    get_ollama_models,
    get_sections,
    stream_generate,
)

_HERE = Path(__file__).parent
if getattr(sys, "frozen", False):
    _env_repo = os.environ.get("CITL_REPO", "").strip()
    REPO = Path(_env_repo) if _env_repo and Path(_env_repo).is_dir() else Path(sys.executable).parent.parent.parent
else:
    REPO = _HERE.parent
DOCS_DIR = REPO / "documents"
APOTHECARY_DIR = REPO / "factbook-assistant" / "fonts" / "doc_composer" / "apothecary"

# ── RTC Navy / Gold dark palette ──────────────────────────────────────────────
C = {
    "bg":        "#0D1B2A",   # deep navy canvas
    "panel":     "#112236",   # card / panel surface
    "panel_alt": "#162B40",   # lifted panel
    "panel_deep":"#0A1520",   # sunken / notebook
    "card_sel":  "#1E4060",   # selected item
    "text":      "#D4E4F5",   # primary text — soft blue-white
    "muted":     "#7A9BBE",   # secondary / labels
    "faint":     "#3E5A78",   # placeholder / disabled
    "accent":    "#3A8FD4",   # RTC blue CTA
    "gold":      "#E89820",   # RTC gold highlights
    "btn":       "#1A3550",   # normal button
    "btn_hi":    "#235272",   # hover button
    "btn_acc":   "#1A4A7A",   # accent button (generate)
    "btn_gold":  "#5A3A00",   # gold action button
    "danger":    "#8B2020",   # stop / destructive
    "notebk":    "#0C1A2C",   # editor background
    "line":      "#1D3050",   # border
}
_F = "Segoe UI" if sys.platform == "win32" else "Ubuntu"

APP_NAME = "CITL Document Composer"
APP_VERSION = "v2.0"

def check_for_updates() -> Optional[str]:
    """
    Check if there's a newer version of this app available in the CITL repo.
    Returns the path to the newer version if found, None otherwise.
    """
    try:
        # Check if we're running from a repo
        if getattr(sys, "frozen", False):
            # Running as exe - check the repo path
            repo_path = os.environ.get("CITL_REPO", "").strip()
            if not repo_path or not Path(repo_path).is_dir():
                return None
            repo = Path(repo_path)
        else:
            # Running as script - check parent directory
            repo = _HERE.parent

        # Look for newer version in repo
        script_path = repo / "factbook-assistant" / "citl_doc_composer.py"
        if not script_path.exists():
            return None

        # Check if the repo version is newer by comparing file modification times
        current_path = Path(__file__)
        if script_path.stat().st_mtime > current_path.stat().st_mtime:
            return str(script_path)

        return None
    except Exception:
        return None

def update_from_repo() -> bool:
    """
    Update this app from the newer version in the CITL repo.
    Returns True if update was successful.
    """
    try:
        newer_path = check_for_updates()
        if not newer_path:
            return False

        # Backup current version
        current_path = Path(__file__)
        backup_path = current_path.with_suffix('.bak')
        if backup_path.exists():
            backup_path.unlink()
        current_path.rename(backup_path)

        # Copy newer version
        import shutil
        shutil.copy2(newer_path, current_path)

        return True
    except Exception:
        return False

_RE_NUM   = re.compile(r"^(\d+(?:\.\d+){0,4})[\).:]?\s+(.+)$")
_RE_ALPHA = re.compile(r"^([a-zA-Z])[\).:]\s+(.+)$")
_RE_BUL   = re.compile(r"^(\s*)[-*\u2022]\s+(.+)$")
_RE_CALL  = re.compile(r"^(TIP|NOTE|WARNING):\s*(.+)$", re.IGNORECASE)
_RE_SHOT  = re.compile(
    r"^(SCREENSHOT(?:\s+PLACEHOLDER)?|IMAGE|\[SCREENSHOT)(?:\s+#?\d+)?\s*[:\]]\s*(.*)$",
    re.IGNORECASE,
)
_RE_SUBHD = re.compile(
    r"^(Menu Path|Expected Result|Validation|Troubleshooting|Context)\s*:\s*(.*)$",
    re.IGNORECASE,
)
_RE_BOLD_ITALIC = re.compile(r"\*\*\*(.+?)\*\*\*")
_RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
_RE_ITALIC = re.compile(r"\*(.+?)\*")


# ── Content rendering helpers ─────────────────────────────────────────────────
def _list_style_name(ordered: bool, level: int) -> str:
    idx = max(1, min(3, level + 1))
    if ordered:
        return "List Number" if idx == 1 else f"List Number {idx}"
    return "List Bullet" if idx == 1 else f"List Bullet {idx}"


def _add_list_line(doc, text: str, ordered: bool, level: int) -> None:
    text = text.strip()
    if not text:
        return
    style = _list_style_name(ordered, level)
    try:
        p = doc.add_paragraph(style=style)
        p.text = text
    except Exception:
        prefix = f"{level + 1}. " if ordered else "- "
        doc.add_paragraph(prefix + text)


def _add_numbered_screenshot(doc, shot_state: dict, caption: str) -> None:
    shot_state["idx"] = int(shot_state.get("idx", 0)) + 1
    number = int(shot_state["idx"])
    caption = caption.strip() or "Capture this step."
    # If a real image is available for this index, insert it
    images: list = shot_state.get("images", [])
    img_idx = number - 1
    if img_idx < len(images):
        try:
            from docx.shared import Inches
            p = doc.add_paragraph()
            p.add_run().add_picture(images[img_idx], width=Inches(5.5))
            from docx.shared import Pt, RGBColor
            cap_p = doc.add_paragraph()
            cap_run = cap_p.add_run(f"Figure {number}: {caption}")
            cap_run.font.size = Pt(9)
            cap_run.font.italic = True
            return
        except Exception:
            pass
    add_screenshot_placeholder(doc, f"Screenshot {number:02d}: {caption}")


def _strip_inline_markup(text: str) -> str:
    """Remove **bold**, *italic*, ***bold-italic*** markers for plain docx runs."""
    text = _RE_BOLD_ITALIC.sub(r"\1", text)
    text = _RE_BOLD.sub(r"\1", text)
    text = _RE_ITALIC.sub(r"\1", text)
    return text


def _render_structured_content(doc, content: str, shot_state: dict) -> None:
    from citl_doc_theme import add_h1_with_bar, add_h3
    try:
        from citl_doc_theme import add_h2
    except ImportError:
        def add_h2(d, t): d.add_heading(t, level=2)  # type: ignore

    para_buf: List[str] = []
    pending_major_step: Optional[str] = None
    pending_major_has_shot = False

    def _flush_paragraph() -> None:
        if para_buf:
            text = _strip_inline_markup(" ".join(para_buf).strip())
            add_body(doc, text)
            para_buf.clear()

    def _flush_pending_major_step() -> None:
        nonlocal pending_major_step, pending_major_has_shot
        if pending_major_step and not pending_major_has_shot:
            _add_numbered_screenshot(doc, shot_state, f"Evidence for step: {pending_major_step}")
        pending_major_step = None
        pending_major_has_shot = False

    for raw in content.splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            _flush_paragraph()
            continue

        # Markdown-style headings
        if stripped.startswith("### "):
            _flush_paragraph()
            add_h3(doc, stripped[4:].strip())
            continue
        if stripped.startswith("## "):
            _flush_paragraph()
            add_h2(doc, stripped[3:].strip())
            continue
        if stripped.startswith("# "):
            _flush_paragraph()
            _flush_pending_major_step()
            add_h1_with_bar(doc, stripped[2:].strip())
            continue

        m = _RE_CALL.match(stripped)
        if m:
            _flush_paragraph()
            add_callout(doc, _strip_inline_markup(m.group(2).strip()), m.group(1).lower())
            continue

        m = _RE_SHOT.match(stripped)
        if m:
            _flush_paragraph()
            _add_numbered_screenshot(doc, shot_state, m.group(2).strip() or "Capture this step.")
            pending_major_has_shot = True
            continue

        m = _RE_SUBHD.match(stripped)
        if m:
            _flush_paragraph()
            add_h3(doc, m.group(1).strip().title())
            if m.group(2).strip():
                add_body(doc, _strip_inline_markup(m.group(2).strip()))
            continue

        m = _RE_NUM.match(stripped)
        if m:
            _flush_paragraph()
            step_depth = m.group(1).count(".")
            if step_depth == 0:
                _flush_pending_major_step()
                pending_major_step = m.group(2).strip()
                pending_major_has_shot = False
            depth = min(2, step_depth)
            _add_list_line(doc, _strip_inline_markup(m.group(2).strip()), ordered=True, level=depth)
            continue

        m = _RE_ALPHA.match(stripped)
        if m:
            _flush_paragraph()
            _add_list_line(doc, _strip_inline_markup(m.group(2).strip()), ordered=True, level=1)
            continue

        m = _RE_BUL.match(line)
        if m:
            _flush_paragraph()
            indent = len(m.group(1).replace("\t", "    "))
            depth = min(2, indent // 2)
            _add_list_line(doc, _strip_inline_markup(m.group(2).strip()), ordered=False, level=depth)
            continue

        para_buf.append(stripped)

    _flush_paragraph()
    _flush_pending_major_step()


def _export_docx(sections: List[dict], meta: dict, out_path: str,
                 embedded_images: Optional[List] = None) -> None:
    from docx import Document
    doc = Document()
    apply_citl_styles(doc, meta.get("doc_style"))
    shot_state: dict = {"idx": 0, "images": embedded_images or []}

    for sec in sections:
        if sec["id"] == "cover":
            add_cover_page(doc, meta)
            continue
        add_h1_with_bar(doc, sec["title"])
        content = sec.get("content", "").strip()
        if not content:
            add_body(doc, f"[{sec['title']} — content not yet generated]")
            add_rule(doc)
            continue
        _render_structured_content(doc, content, shot_state)
        add_rule(doc)

    doc.save(out_path)


def _export_pdf(docx_path: str, pdf_path: str, log=print) -> bool:
    """Try docx2pdf → LibreOffice → fail with instructions."""
    # Method 1: docx2pdf
    try:
        from docx2pdf import convert  # type: ignore
        convert(docx_path, pdf_path)
        if Path(pdf_path).exists():
            return True
    except ImportError:
        pass
    except Exception as exc:
        log(f"docx2pdf: {exc}")

    # Method 2: LibreOffice / soffice
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            out_dir = str(Path(pdf_path).parent)
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
                capture_output=True, timeout=90,
            )
            # LibreOffice names the PDF after the input stem
            generated = Path(out_dir) / (Path(docx_path).stem + ".pdf")
            if generated.exists():
                if str(generated) != pdf_path:
                    generated.replace(pdf_path)
                return True
        except Exception as exc:
            log(f"LibreOffice: {exc}")

    return False


# ── Font installation (full apothecary recursive scan) ───────────────────────
_FONT_EXTS = {".ttf", ".otf", ".ttc"}


def _install_all_apothecary_fonts(log=print) -> dict:
    """
    Recursively scan factbook-assistant/fonts/doc_composer/apothecary/
    and install all font files to the user-level Windows font directory.
    No administrator rights required on Windows 10+.
    """
    if sys.platform != "win32":
        log("[SKIP] Font install is Windows-only.")
        return {}

    src_dir = APOTHECARY_DIR
    if not src_dir.exists():
        log(f"[MISS] Font directory not found: {src_dir}")
        return {}

    font_dest = Path(os.environ.get("LOCALAPPDATA", "")) / "Microsoft" / "Windows" / "Fonts"
    font_dest.mkdir(parents=True, exist_ok=True)

    results: dict = {}
    installed = 0
    skipped = 0
    errors = 0

    all_fonts = [p for p in src_dir.rglob("*") if p.is_file() and p.suffix.lower() in _FONT_EXTS]
    log(f"Found {len(all_fonts)} font files in {src_dir}")

    for font_file in all_fonts:
        dst = font_dest / font_file.name
        key = font_file.name
        try:
            if not dst.exists():
                shutil.copy2(font_file, dst)
                installed += 1
            else:
                skipped += 1

            with winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts",
                access=winreg.KEY_SET_VALUE,
            ) as k:
                kind = "OpenType" if font_file.suffix.lower() == ".otf" else "TrueType"
                winreg.SetValueEx(k, f"{font_file.stem} ({kind})", 0, winreg.REG_SZ, str(dst))

            ctypes.windll.gdi32.AddFontResourceW(str(dst))
            results[key] = True
        except Exception as exc:
            results[key] = False
            errors += 1
            log(f"[ERR] {key}: {exc}")

    if installed + skipped > 0:
        try:
            ctypes.windll.user32.SendMessageW(0xFFFF, 0x001D, 0, 0)
        except Exception:
            pass

    log(f"Font install complete: {installed} new, {skipped} already present, {errors} errors.")
    return results


# ── SnipOverlay — full-screen screenshot selector ────────────────────────────
class SnipOverlay(tk.Toplevel):
    """
    Full-screen screen capture overlay.
    Shows the captured desktop screenshot; user drags to select a region.
    Calls on_done(PIL.Image | None) when finished.
    """
    def __init__(self, master: tk.Misc, bg_screenshot, on_done):
        super().__init__(master)
        self._on_done = on_done
        self._bg_img = bg_screenshot
        self._x0 = self._y0 = 0
        self._rect_id = None
        self._cancelled = False

        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()

        self.overrideredirect(True)
        self.geometry(f"{sw}x{sh}+0+0")
        self.attributes("-topmost", True)
        self.configure(bg="black", cursor="crosshair")
        self.lift()
        self.focus_force()

        # Scale screenshot to screen size and display as canvas background
        from PIL import ImageTk
        scaled = bg_screenshot.resize((sw, sh))
        self._tk_bg = ImageTk.PhotoImage(scaled)

        self._canvas = tk.Canvas(
            self, width=sw, height=sh,
            bg="black", cursor="crosshair",
            highlightthickness=0,
        )
        self._canvas.pack(fill="both", expand=True)
        self._canvas.create_image(0, 0, anchor="nw", image=self._tk_bg, tags="bg")
        # Dim overlay
        self._canvas.create_rectangle(0, 0, sw, sh,
                                      fill="navy", stipple="gray25",
                                      outline="", tags="dim")

        # Instruction label
        self._canvas.create_rectangle(sw // 2 - 230, 10, sw // 2 + 230, 44,
                                      fill="#0D1B2A", outline="#3A8FD4", width=1)
        self._canvas.create_text(sw // 2, 27,
                                 text="Drag to select area  •  Esc to cancel",
                                 fill="#D4E4F5", font=("Segoe UI", 13))

        self._canvas.bind("<ButtonPress-1>", self._press)
        self._canvas.bind("<B1-Motion>", self._drag)
        self._canvas.bind("<ButtonRelease-1>", self._release)
        self.bind("<Escape>", lambda _: self._finish(None))
        self.bind("<FocusOut>", lambda _: self._finish(None))

    def _press(self, e):
        self._x0, self._y0 = e.x, e.y
        if self._rect_id:
            self._canvas.delete(self._rect_id)
        self._rect_id = self._canvas.create_rectangle(
            e.x, e.y, e.x, e.y,
            outline="#3A8FD4", width=2, dash=(5, 3), tags="sel",
        )

    def _drag(self, e):
        if self._rect_id:
            self._canvas.coords(self._rect_id, self._x0, self._y0, e.x, e.y)

    def _release(self, e):
        x0, y0 = min(self._x0, e.x), min(self._y0, e.y)
        x1, y1 = max(self._x0, e.x), max(self._y0, e.y)
        if x1 - x0 < 5 or y1 - y0 < 5:
            self._finish(None)
            return
        # Scale coords back to actual screenshot dimensions
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        orig_w, orig_h = self._bg_img.size
        scale_x = orig_w / sw
        scale_y = orig_h / sh
        rx0 = int(x0 * scale_x)
        ry0 = int(y0 * scale_y)
        rx1 = int(x1 * scale_x)
        ry1 = int(y1 * scale_y)
        cropped = self._bg_img.crop((rx0, ry0, rx1, ry1))
        self._finish(cropped)

    def _finish(self, img):
        if self._cancelled:
            return
        self._cancelled = True
        try:
            self.destroy()
        except Exception:
            pass
        self._on_done(img)


# ── Screenshot tray widget ────────────────────────────────────────────────────
class ScreenshotTray(tk.Frame):
    """
    Horizontal scrollable tray showing thumbnail previews of captured screenshots.
    Click a thumbnail to insert a [SCREENSHOT: N] marker at the editor cursor.
    Right-click (or click ×) to remove.
    """
    THUMB_W = 80
    THUMB_H = 60

    def __init__(self, parent, on_insert, **kw):
        kw.setdefault("bg", C["panel_deep"])
        kw.setdefault("height", self.THUMB_H + 28)
        super().__init__(parent, **kw)
        self._on_insert = on_insert  # callback(index: int, img: PIL.Image)
        self._images: list = []      # PIL Images
        self._tk_thumbs: list = []   # PhotoImage refs (prevent GC)

        self._canvas = tk.Canvas(self, bg=C["panel_deep"], height=self.THUMB_H + 26,
                                 highlightthickness=0, bd=0)
        self._scroll = ttk.Scrollbar(self, orient="horizontal",
                                     command=self._canvas.xview)
        self._canvas.configure(xscrollcommand=self._scroll.set)
        self._scroll.pack(side="bottom", fill="x")
        self._canvas.pack(side="top", fill="both", expand=True)

        self._frame = tk.Frame(self._canvas, bg=C["panel_deep"])
        self._canvas.create_window(0, 0, anchor="nw", window=self._frame)
        self._frame.bind("<Configure>", self._on_frame_resize)

        self._placeholder = tk.Label(
            self._frame,
            text="No screenshots — click  📷 Snip  or browse to attach",
            font=(_F, 8, "italic"), bg=C["panel_deep"], fg=C["faint"],
        )
        self._placeholder.pack(pady=8)

    def _on_frame_resize(self, _e):
        self._canvas.configure(scrollregion=self._canvas.bbox("all"))

    def add_image(self, img) -> int:
        """Add a PIL Image to the tray. Returns index."""
        from PIL import ImageTk
        self._images.append(img)
        idx = len(self._images) - 1

        if self._placeholder.winfo_exists():
            try:
                self._placeholder.pack_forget()
            except Exception:
                pass

        cell = tk.Frame(self._frame, bg=C["panel"], relief="flat",
                        padx=2, pady=2)
        cell.pack(side="left", padx=4, pady=3)

        thumb = img.copy()
        thumb.thumbnail((self.THUMB_W, self.THUMB_H))
        tk_img = ImageTk.PhotoImage(thumb)
        self._tk_thumbs.append(tk_img)

        lbl = tk.Label(cell, image=tk_img, bg=C["line"], cursor="hand2",
                       relief="flat")
        lbl.image = tk_img  # extra ref
        lbl.pack()

        num_lbl = tk.Label(cell, text=f"#{idx + 1}", font=(_F, 7),
                           bg=C["panel"], fg=C["muted"])
        num_lbl.pack()

        remove_btn = tk.Button(
            cell, text="×", font=(_F, 8, "bold"),
            bg=C["danger"], fg=C["text"], relief="flat", bd=0,
            padx=3, pady=0, cursor="hand2",
            command=lambda i=idx, c=cell: self._remove(i, c),
        )
        remove_btn.pack()

        lbl.bind("<Button-1>", lambda _e, i=idx: self._on_insert(i, self._images[i]))
        lbl.bind("<Button-3>", lambda _e, i=idx, c=cell: self._remove(i, c))

        self._canvas.configure(scrollregion=self._canvas.bbox("all"))
        return idx

    def _remove(self, idx: int, cell: tk.Frame):
        try:
            cell.destroy()
        except Exception:
            pass
        if idx < len(self._images):
            self._images[idx] = None  # tombstone
        if not any(img is not None for img in self._images):
            self._placeholder.pack(pady=8)

    def get_live_images(self) -> list:
        """Return all non-removed PIL Images in order (None entries excluded)."""
        return [img for img in self._images if img is not None]

    def count(self) -> int:
        return sum(1 for img in self._images if img is not None)


# ── DocComposer main application ──────────────────────────────────────────────
class DocComposer:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title(f"{APP_NAME}  {APP_VERSION}")
        self.root.configure(bg=C["bg"])
        self.root.minsize(1100, 720)

        self._sections: List[dict] = []
        self._current_idx: int = -1        # -1 = free-compose mode
        self._generating = False
        self._gen_stop = threading.Event()

        self._models: List[dict] = []
        self._model_var = tk.StringVar()
        self._model_disp = tk.StringVar(value="Checking Ollama…")

        self._template_var  = tk.StringVar(value=TEMPLATE_NAMES[0])
        self._app_name_var  = tk.StringVar(value="CITL Application")
        self._version_var   = tk.StringVar(value="v1.0")
        self._author_var    = tk.StringVar(value="CITL Staff")
        self._date_var      = tk.StringVar(value=str(date.today()))
        self._subtitle_var  = tk.StringVar()
        self._ui_goal_var   = tk.StringVar()

        style_names  = get_doc_style_names()
        active_style = get_active_doc_style()
        if active_style not in style_names and style_names:
            active_style = style_names[0]
        self._doc_style_var = tk.StringVar(value=active_style)

        self._status_var   = tk.StringVar(value="Ready — compose freely or select a template section")
        self._font_status  = tk.StringVar(value="Checking fonts…")
        self._warned_non_vision = False

        DOCS_DIR.mkdir(parents=True, exist_ok=True)

        self._build_ui()
        self.root.after(300, self._detect_models)
        self.root.after(500, self._check_fonts)

    # ── UI construction ──────────────────────────────────────────────────────

    def _build_ui(self) -> None:
        self._build_header()
        self._build_toolbar()
        self._build_status_bar()

        body = tk.Frame(self.root, bg=C["bg"])
        body.pack(fill="both", expand=True, padx=6, pady=(4, 6))
        body.columnconfigure(0, weight=0, minsize=200)
        body.columnconfigure(1, weight=1)
        body.columnconfigure(2, weight=0, minsize=290)
        body.rowconfigure(0, weight=1)

        self._build_left(body)
        self._build_center(body)
        self._build_right(body)

    def _build_header(self) -> None:
        hdr = tk.Frame(self.root, bg=C["panel"])
        hdr.pack(fill="x")
        tk.Frame(hdr, bg=C["gold"], height=3).pack(fill="x")
        hi = tk.Frame(hdr, bg=C["panel"])
        hi.pack(fill="x", padx=16, pady=8)
        tk.Label(hi, text=APP_NAME, font=(_F, 17, "bold"),
                 bg=C["panel"], fg=C["text"]).pack(side="left")
        tk.Label(hi, text=APP_VERSION, font=(_F, 9, "bold"),
                 bg=C["panel"], fg=C["gold"]).pack(side="left", padx=6)
        tk.Label(hi, textvariable=self._model_disp, font=(_F, 9, "italic"),
                 bg=C["panel"], fg=C["muted"]).pack(side="right")

    def _build_toolbar(self) -> None:
        tb = tk.Frame(self.root, bg=C["panel_alt"], pady=4)
        tb.pack(fill="x")

        def _tbtn(text, cmd, gold=False, tip=""):
            b = tk.Button(
                tb, text=text, font=(_F, 9), relief="flat", bd=0,
                bg=C["btn_gold"] if gold else C["btn"],
                fg=C["gold"] if gold else C["text"],
                activebackground=C["btn_hi"],
                padx=7, pady=3, cursor="hand2", command=cmd,
            )
            b.pack(side="left", padx=2)
            return b

        # Format insertion buttons
        _tbtn("Title",   lambda: self._insert_fmt("# "))
        _tbtn("H1",      lambda: self._insert_fmt("## "))
        _tbtn("H2",      lambda: self._insert_fmt("### "))
        tk.Frame(tb, bg=C["line"], width=1).pack(side="left", fill="y", padx=4)
        _tbtn("B",       lambda: self._wrap_fmt("**", "**"))
        _tbtn("I",       lambda: self._wrap_fmt("*", "*"))
        tk.Frame(tb, bg=C["line"], width=1).pack(side="left", fill="y", padx=4)
        _tbtn("1.",      lambda: self._insert_fmt("1. "))
        _tbtn("1.1.",    lambda: self._insert_fmt("1.1. "))
        _tbtn("•",       lambda: self._insert_fmt("- "))
        tk.Frame(tb, bg=C["line"], width=1).pack(side="left", fill="y", padx=4)
        _tbtn("TIP:",    lambda: self._insert_fmt("TIP: "))
        _tbtn("NOTE:",   lambda: self._insert_fmt("NOTE: "))
        _tbtn("WARN:",   lambda: self._insert_fmt("WARNING: "))
        tk.Frame(tb, bg=C["line"], width=1).pack(side="left", fill="y", padx=4)
        _tbtn("[IMG]",   self._insert_screenshot_marker)
        tk.Frame(tb, bg=C["bg"], width=1).pack(side="left", fill="y", padx=2, expand=True)
        _tbtn("📷 Snip", self._launch_snip, gold=True)

    def _build_status_bar(self) -> None:
        sb = tk.Frame(self.root, bg=C["panel_deep"])
        sb.pack(fill="x")
        tk.Label(sb, textvariable=self._status_var, font=(_F, 8),
                 bg=C["panel_deep"], fg=C["accent"]).pack(side="left", padx=10)
        tk.Label(sb, textvariable=self._font_status, font=(_F, 8),
                 bg=C["panel_deep"], fg=C["muted"]).pack(side="right", padx=10)

    def _build_left(self, parent) -> None:
        left = tk.Frame(parent, bg=C["panel"], relief="flat", bd=0)
        left.grid(row=0, column=0, sticky="nsew", padx=(0, 5))
        left.rowconfigure(2, weight=1)

        self._lbl(left, "TEMPLATE").pack(anchor="w", padx=8, pady=(8, 2))
        self._tmpl_combo = ttk.Combobox(
            left, textvariable=self._template_var,
            values=TEMPLATE_NAMES, state="readonly", font=(_F, 9), width=24,
        )
        self._tmpl_combo.pack(fill="x", padx=8, pady=(0, 6))
        self._tmpl_combo.bind("<<ComboboxSelected>>", self._on_template_change)

        self._lbl(left, "SECTIONS").pack(anchor="w", padx=8, pady=(4, 2))
        self._section_lb = tk.Listbox(
            left, font=(_F, 9), bg=C["notebk"], fg=C["text"],
            selectbackground=C["card_sel"], selectforeground=C["gold"],
            activestyle="none", bd=0, highlightthickness=0, relief="flat",
            selectmode="single",
        )
        self._section_lb.pack(fill="both", expand=True, padx=8, pady=(0, 4))
        self._section_lb.bind("<<ListboxSelect>>", self._on_section_select)

        # Free compose button
        self._free_btn = self._btn(left, "Free Compose", C["btn"], self._free_compose_mode)
        self._free_btn.pack(fill="x", padx=8, pady=(0, 2))

        self._gen_all_btn = self._btn(left, "⚡ Generate All", C["btn_acc"], self._generate_all)
        self._gen_all_btn.pack(fill="x", padx=8, pady=(0, 2))

        self._stop_btn = self._btn(left, "■ Stop", C["danger"], self._stop_generate)
        self._stop_btn.pack(fill="x", padx=8, pady=(0, 8))
        self._stop_btn.config(state="disabled")

    def _build_center(self, parent) -> None:
        center = tk.Frame(parent, bg=C["bg"])
        center.grid(row=0, column=1, sticky="nsew", padx=4)
        center.rowconfigure(1, weight=1)
        center.columnconfigure(0, weight=1)

        self._sec_label = tk.Label(
            center, text="Free Compose", font=(_F, 11, "bold"),
            bg=C["bg"], fg=C["gold"], anchor="w",
        )
        self._sec_label.grid(row=0, column=0, sticky="ew", pady=(0, 3))

        self._editor = scrolledtext.ScrolledText(
            center, font=("Georgia", 10),
            bg=C["notebk"], fg="#D4E4F5",
            insertbackground=C["text"],
            selectbackground=C["card_sel"],
            relief="flat", wrap="word",
            state="normal",           # ALWAYS ACTIVE
            undo=True,
        )
        self._editor.grid(row=1, column=0, sticky="nsew")
        self._editor.bind("<KeyRelease>", self._on_editor_change)
        self._editor.focus_set()

        # Screenshot tray
        tray_lbl = tk.Label(center, text="SCREENSHOT TRAY",
                            font=(_F, 7, "bold"), bg=C["bg"], fg=C["muted"], anchor="w")
        tray_lbl.grid(row=2, column=0, sticky="ew", pady=(6, 1))
        self._tray = ScreenshotTray(
            center,
            on_insert=self._on_tray_insert,
        )
        self._tray.grid(row=3, column=0, sticky="ew")

        # Action row
        ec = tk.Frame(center, bg=C["bg"])
        ec.grid(row=4, column=0, sticky="ew", pady=(6, 0))
        self._gen_sec_btn = self._btn(ec, "⚡ Generate Section", C["btn_acc"], self._generate_section_cmd)
        self._gen_sec_btn.pack(side="left", padx=(0, 5))
        self._btn(ec, "Clear", C["btn"], self._clear_section).pack(side="left", padx=(0, 5))
        tk.Label(ec, text="Ctrl+Z to undo  •  edits auto-saved",
                 font=(_F, 8, "italic"), bg=C["bg"], fg=C["faint"]).pack(side="right")

    def _build_right(self, parent) -> None:
        right = tk.Frame(parent, bg=C["panel"], relief="flat", bd=0)
        right.grid(row=0, column=2, sticky="nsew")
        right.columnconfigure(1, weight=1)

        row = 0

        def _row(label, var):
            nonlocal row
            tk.Label(right, text=label, font=(_F, 8), bg=C["panel"],
                     fg=C["muted"], anchor="w").grid(
                row=row, column=0, sticky="ew", padx=(8, 4), pady=2)
            tk.Entry(right, textvariable=var, font=(_F, 9), bg=C["notebk"],
                     fg=C["text"], insertbackground=C["text"], relief="flat",
                     highlightthickness=0,
                     ).grid(row=row, column=1, sticky="ew", padx=(0, 8), pady=2)
            row += 1

        self._lbl(right, "DOCUMENT INFO").grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(8, 4))
        row += 1
        _row("App name:",  self._app_name_var)
        _row("Subtitle:",  self._subtitle_var)
        _row("Author:",    self._author_var)
        _row("Version:",   self._version_var)
        _row("Date:",      self._date_var)
        _row("UI goal:",   self._ui_goal_var)

        self._lbl(right, "AI TOPIC / PROMPT").grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(8, 2))
        row += 1
        self._topic_text = tk.Text(
            right, font=(_F, 9), bg=C["notebk"], fg=C["text"],
            insertbackground=C["text"], relief="flat", height=4, wrap="word",
            highlightthickness=0,
        )
        self._topic_text.grid(row=row, column=0, columnspan=2,
                               sticky="ew", padx=8, pady=(0, 6))
        row += 1

        self._lbl(right, "OLLAMA MODEL").grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(4, 2))
        row += 1
        self._model_combo = ttk.Combobox(
            right, textvariable=self._model_var, font=(_F, 9),
            state="readonly", width=28,
        )
        self._model_combo.grid(row=row, column=0, columnspan=2,
                                sticky="ew", padx=8, pady=(0, 2))
        row += 1
        self._btn(right, "Refresh Models", C["btn"], self._detect_models).grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(0, 6))
        row += 1

        self._lbl(right, "TYPOGRAPHY PRESET").grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(4, 2))
        row += 1
        self._style_combo = ttk.Combobox(
            right, textvariable=self._doc_style_var,
            values=get_doc_style_names(), state="readonly",
            font=(_F, 9), width=28,
        )
        self._style_combo.grid(row=row, column=0, columnspan=2,
                                sticky="ew", padx=8, pady=(0, 2))
        self._style_combo.bind("<<ComboboxSelected>>", self._on_doc_style_change)
        row += 1

        self._btn(right, "Install Apothecary Fonts", C["btn"], self._install_fonts).grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(0, 8))
        row += 1

        tk.Frame(right, bg=C["line"], height=1).grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8)
        row += 1

        self._lbl(right, "EXPORT").grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(8, 4))
        row += 1
        self._export_docx_btn = self._btn(right, "Export  .docx", C["btn_acc"], self._export_docx_cmd)
        self._export_docx_btn.grid(row=row, column=0, columnspan=2,
                                   sticky="ew", padx=8, pady=(0, 4))
        row += 1
        self._export_pdf_btn = self._btn(right, "Export  PDF", C["btn_gold"], self._export_pdf_cmd)
        self._export_pdf_btn.grid(row=row, column=0, columnspan=2,
                                  sticky="ew", padx=8, pady=(0, 4))
        row += 1
        self._btn(right, "Open Documents Folder", C["btn"], self._open_docs_dir).grid(
            row=row, column=0, columnspan=2, sticky="ew", padx=8, pady=(0, 8))

    # ── Widget helpers ────────────────────────────────────────────────────────

    def _lbl(self, parent, text: str) -> tk.Label:
        return tk.Label(parent, text=text, font=(_F, 8, "bold"),
                        bg=C["panel"], fg=C["accent"])

    def _btn(self, parent, text: str, bg: str, cmd) -> tk.Button:
        return tk.Button(
            parent, text=text, font=(_F, 9), bg=bg,
            fg=C["text"], activebackground=C["btn_hi"],
            relief="flat", bd=0, padx=8, pady=5, cursor="hand2",
            command=cmd,
        )

    # ── Formatting toolbar actions ────────────────────────────────────────────

    def _insert_fmt(self, prefix: str) -> None:
        """Insert prefix at start of current line or at cursor."""
        try:
            pos = self._editor.index("insert")
            line_start = f"{pos.split('.')[0]}.0"
            self._editor.insert(line_start, prefix)
            self._editor.mark_set("insert", f"{line_start}+{len(prefix)}c")
            self._editor.focus_set()
        except Exception:
            pass

    def _wrap_fmt(self, before: str, after: str) -> None:
        """Wrap selected text with before/after markers, or insert at cursor."""
        try:
            try:
                sel = self._editor.get("sel.first", "sel.last")
                self._editor.delete("sel.first", "sel.last")
                self._editor.insert("insert", f"{before}{sel}{after}")
            except tk.TclError:
                self._editor.insert("insert", f"{before}{after}")
                pos = self._editor.index("insert")
                line, col = pos.split(".")
                new_col = max(0, int(col) - len(after))
                self._editor.mark_set("insert", f"{line}.{new_col}")
            self._editor.focus_set()
        except Exception:
            pass

    def _insert_screenshot_marker(self) -> None:
        count = self._tray.count() + 1
        self._editor.insert("insert", f"\n[SCREENSHOT: Figure {count} caption here]\n")
        self._editor.focus_set()

    # ── Snip tool ─────────────────────────────────────────────────────────────

    def _launch_snip(self) -> None:
        try:
            from PIL import ImageGrab  # type: ignore
        except ImportError:
            messagebox.showerror(APP_NAME, "Pillow is required for screen capture.\n  pip install Pillow")
            return
        self.root.withdraw()
        self.root.update()
        self.root.after(180, self._do_capture)

    def _do_capture(self) -> None:
        try:
            from PIL import ImageGrab
            screenshot = ImageGrab.grab()
        except Exception as exc:
            self.root.deiconify()
            messagebox.showerror(APP_NAME, f"Screen capture failed:\n{exc}")
            return

        def _done(img):
            self.root.deiconify()
            if img is not None:
                idx = self._tray.add_image(img)
                self._status_var.set(f"Screenshot #{idx + 1} added to tray — click thumbnail to insert marker")

        try:
            SnipOverlay(self.root, screenshot, _done)
        except Exception as exc:
            self.root.deiconify()
            messagebox.showerror(APP_NAME, f"Snip overlay failed:\n{exc}")

    def _on_tray_insert(self, idx: int, img) -> None:
        """Called when user clicks a tray thumbnail — insert marker in editor."""
        marker = f"\n[SCREENSHOT: Figure {idx + 1} — add caption here]\n"
        self._editor.insert("insert", marker)
        self._editor.focus_set()
        self._status_var.set(f"Screenshot #{idx + 1} marker inserted")

    # ── Section management ────────────────────────────────────────────────────

    def _free_compose_mode(self) -> None:
        self._save_current_section()
        self._current_idx = -1
        self._section_lb.selection_clear(0, "end")
        self._sec_label.config(text="Free Compose")
        self._editor.config(state="normal")
        self._status_var.set("Free compose — write anything, use toolbar for formatting")
        self._editor.focus_set()

    def _on_template_change(self, *_):
        self._save_current_section()
        self._sections = get_sections(self._template_var.get())
        self._current_idx = -1
        self._section_lb.delete(0, "end")
        for s in self._sections:
            self._section_lb.insert("end", s["title"])

    def _on_section_select(self, *_):
        sel = self._section_lb.curselection()
        if not sel:
            return
        idx = int(sel[0])
        self._save_current_section()
        self._load_section(idx)

    def _load_section(self, idx: int) -> None:
        self._current_idx = idx
        sec = self._sections[idx]
        self._sec_label.config(text=sec["title"])
        self._editor.config(state="normal")
        self._editor.delete("1.0", "end")
        if sec["id"] == "cover":
            self._editor.insert("1.0", "(Cover page is auto-generated from the Document Info fields →)")
            self._editor.config(state="disabled")
        else:
            self._editor.insert("1.0", sec.get("content", ""))
            self._editor.config(state="normal")
        self._editor.focus_set()

    def _save_current_section(self) -> None:
        if 0 <= self._current_idx < len(self._sections):
            sec = self._sections[self._current_idx]
            if sec["id"] != "cover":
                sec["content"] = self._editor.get("1.0", "end-1c")

    def _on_editor_change(self, *_):
        self._save_current_section()

    def _clear_section(self) -> None:
        if 0 <= self._current_idx < len(self._sections):
            self._sections[self._current_idx]["content"] = ""
        self._editor.config(state="normal")
        self._editor.delete("1.0", "end")
        self._editor.focus_set()

    # ── Ollama model detection ────────────────────────────────────────────────

    def _detect_models(self) -> None:
        def _run():
            models = get_ollama_models()
            def _upd():
                self._models = models
                if models:
                    names = [m["display"] for m in models]
                    self._model_combo["values"] = names
                    self._model_combo.current(0)
                    self._model_var.set(names[0])
                    best = models[0]
                    vtag = " [vision]" if best.get("is_vision") else ""
                    self._model_disp.set(f"Best: {best['name']}{vtag} ({best['params']}B)")
                else:
                    self._model_disp.set("No Ollama models found")
            self.root.after(0, _upd)
        threading.Thread(target=_run, daemon=True).start()

    def _selected_model_entry(self) -> Optional[dict]:
        idx = self._model_combo.current()
        if 0 <= idx < len(self._models):
            return self._models[idx]
        pick = self._model_var.get().strip()
        for m in self._models:
            if pick in (m.get("display", ""), m.get("name", "")):
                return m
        return self._models[0] if self._models else None

    def _pick_model(self) -> tuple:
        chosen = self._selected_model_entry()
        model_name = chosen["name"] if chosen else get_best_model()
        images = self._tray.get_live_images()
        use_images = bool(images)
        if not model_name:
            return None, False
        if use_images and not (chosen and chosen.get("is_vision")):
            vision = get_best_vision_model()
            if vision:
                model_name = vision
            else:
                if not self._warned_non_vision:
                    messagebox.showwarning(APP_NAME,
                        "No vision-capable Ollama model found.\n"
                        "Screenshots will be omitted from AI generation.")
                    self._warned_non_vision = True
                use_images = False
        return model_name, use_images

    # ── AI generation ─────────────────────────────────────────────────────────

    def _generate_section_cmd(self) -> None:
        self._generate_section(idx=None)

    def _generate_section(self, idx: Optional[int] = None, done_cb=None) -> None:
        if self._generating:
            return
        target = idx if idx is not None else self._current_idx
        if not (0 <= target < len(self._sections)):
            self._status_var.set("Select a template section first (or use Free Compose).")
            return
        sec = self._sections[target]
        if sec["id"] == "cover":
            if done_cb:
                done_cb()
            return

        model_name, use_images = self._pick_model()
        if not model_name:
            messagebox.showwarning(APP_NAME, "No Ollama model found. Make sure Ollama is running.")
            return

        if target != self._current_idx:
            self._save_current_section()
            self._load_section(target)
            self._section_lb.selection_clear(0, "end")
            self._section_lb.selection_set(target)

        self._editor.config(state="normal")
        self._editor.delete("1.0", "end")
        self._sec_label.config(text=f"Generating: {sec['title']}…")
        self._status_var.set(f"Generating '{sec['title']}' with {model_name}…")
        self._set_generating(True)
        self._gen_stop.clear()

        images = self._tray.get_live_images()
        image_io_list: List = []
        if use_images and images:
            for pil_img in images:
                if pil_img is not None:
                    buf = io.BytesIO()
                    pil_img.save(buf, format="PNG")
                    buf.seek(0)
                    image_io_list.append(buf)

        q = stream_generate(
            model=model_name,
            section_prompt=sec["prompt"],
            meta=self._get_meta(),
            token_cb=lambda _t: None,
            done_cb=lambda _ok, _msg: None,
            image_paths=[str(p) for p in self._tray._images if p is not None] if use_images else None,
            system_override=sec.get("system_prompt"),
        )
        self._poll_gen(q, target, done_cb)

    def _poll_gen(self, q: queue.Queue, target: int, done_cb=None) -> None:
        if self._gen_stop.is_set():
            self._finish_gen(target, done_cb)
            return
        try:
            while True:
                msg_type, content = q.get_nowait()
                if msg_type == "token":
                    self._editor.insert("end", content)
                    self._editor.see("end")
                elif msg_type == "done":
                    self._finish_gen(target, done_cb)
                    return
                elif msg_type == "error":
                    self._status_var.set(f"Error: {content}")
                    self._finish_gen(target, done_cb)
                    return
        except queue.Empty:
            pass
        self.root.after(40, lambda: self._poll_gen(q, target, done_cb))

    def _finish_gen(self, target: int, done_cb=None) -> None:
        self._save_current_section()
        sec = self._sections[target]
        self._sec_label.config(text=sec["title"])
        self._status_var.set(f"Done: {sec['title']}")
        self._set_generating(False)
        if done_cb:
            done_cb()

    def _generate_all(self) -> None:
        if self._generating:
            return
        idxs = [i for i, s in enumerate(self._sections) if s["id"] != "cover" and not s.get("content", "").strip()]
        if not idxs:
            idxs = [i for i, s in enumerate(self._sections) if s["id"] != "cover"]
        self._gen_stop.clear()
        self._run_gen_sequence(idxs, 0)

    def _run_gen_sequence(self, idxs: List[int], pos: int) -> None:
        if pos >= len(idxs) or self._gen_stop.is_set():
            msg = "All sections generated." if pos >= len(idxs) else "Generation stopped."
            self._status_var.set(msg)
            self._set_generating(False)
            return
        self._generate_section(
            idx=idxs[pos],
            done_cb=lambda: self.root.after(200, lambda: self._run_gen_sequence(idxs, pos + 1)),
        )

    def _stop_generate(self) -> None:
        self._gen_stop.set()
        self._status_var.set("Stopping…")

    def _set_generating(self, val: bool) -> None:
        self._generating = val
        state = "disabled" if val else "normal"
        self._gen_all_btn.config(state=state)
        self._gen_sec_btn.config(state=state)
        self._stop_btn.config(state="normal" if val else "disabled")
        self._export_docx_btn.config(state=state)
        self._export_pdf_btn.config(state=state)

    # ── Doc style / font checks ───────────────────────────────────────────────

    def _on_doc_style_change(self, *_) -> None:
        set_doc_style(self._doc_style_var.get())
        self._check_fonts()

    def _check_fonts(self) -> None:
        missing = get_missing_fonts(self._doc_style_var.get())
        payload = get_doc_style_fonts(self._doc_style_var.get())
        if not missing:
            self._font_status.set(
                f"✓ {payload['body']}  ·  {payload['heading']}  ·  {payload['caption']}"
            )
        else:
            self._font_status.set(f"Missing: {', '.join(missing)} — click Install")

    def _install_fonts(self) -> None:
        self._font_status.set("Installing fonts…")
        self._status_var.set("Scanning Apothecary font library…")

        def _run():
            log_lines: List[str] = []
            results = _install_all_apothecary_fonts(log=log_lines.append)
            ok    = sum(1 for v in results.values() if v)
            fail  = sum(1 for v in results.values() if not v)
            total = len(results)

            def _upd():
                if total == 0:
                    self._font_status.set("Font directory not found — check repo structure.")
                    self._status_var.set("Font install: directory missing.")
                else:
                    self._font_status.set(
                        f"Installed {ok}/{total} fonts" + (f" ({fail} errors)" if fail else " ✓")
                    )
                    self._status_var.set(
                        f"Font install complete: {ok} installed, {fail} errors."
                    )
                self._check_fonts()

            self.root.after(0, _upd)

        threading.Thread(target=_run, daemon=True).start()

    # ── Metadata ──────────────────────────────────────────────────────────────

    def _get_meta(self) -> dict:
        topic = self._topic_text.get("1.0", "end").strip()
        return {
            "app_name":      self._app_name_var.get(),
            "version":       self._version_var.get(),
            "author":        self._author_var.get(),
            "date":          self._date_var.get(),
            "subtitle":      self._subtitle_var.get(),
            "topic":         topic,
            "ui_goal":       self._ui_goal_var.get(),
            "template_type": self._template_var.get(),
            "title":         self._app_name_var.get(),
            "doc_style":     self._doc_style_var.get(),
        }

    # ── Export ────────────────────────────────────────────────────────────────

    def _safe_filename(self, base: str, ext: str) -> str:
        safe = "".join(c if c.isalnum() or c in " _-" else "_" for c in base)
        return f"CITL_{safe}_{self._template_var.get().replace(' ', '_')}{ext}"

    def _export_docx_cmd(self) -> None:
        self._save_current_section()
        meta = self._get_meta()
        set_doc_style(meta.get("doc_style"))
        default_name = self._safe_filename(meta["app_name"], ".docx")
        out_path = filedialog.asksaveasfilename(
            title="Save CITL Document (.docx)",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All files", "*.*")],
            initialdir=str(DOCS_DIR),
            initialfile=default_name,
        )
        if not out_path:
            return
        self._do_export_docx(out_path, meta, open_after=True)

    def _do_export_docx(self, out_path: str, meta: dict, open_after: bool = False) -> bool:
        self._status_var.set("Exporting .docx…")
        self.root.update_idletasks()
        images = self._tray.get_live_images()
        img_streams = []
        for pil_img in images:
            buf = io.BytesIO()
            pil_img.save(buf, format="PNG")
            buf.seek(0)
            img_streams.append(buf)
        try:
            _export_docx(self._sections, meta, out_path, embedded_images=img_streams)
            self._status_var.set(f"Saved: {Path(out_path).name}")
            if open_after and messagebox.askyesno(APP_NAME, f"Saved:\n{out_path}\n\nOpen now?"):
                self._open_path(Path(out_path))
            return True
        except Exception as exc:
            self._status_var.set(f"Export failed: {exc}")
            messagebox.showerror(APP_NAME,
                f"Export failed:\n{exc}\n\nMake sure python-docx is installed:\n  pip install python-docx")
            return False

    def _export_pdf_cmd(self) -> None:
        self._save_current_section()
        meta = self._get_meta()
        set_doc_style(meta.get("doc_style"))
        default_name = self._safe_filename(meta["app_name"], ".pdf")
        out_path = filedialog.asksaveasfilename(
            title="Save CITL Document (PDF)",
            defaultextension=".pdf",
            filetypes=[("PDF Document", "*.pdf"), ("All files", "*.*")],
            initialdir=str(DOCS_DIR),
            initialfile=default_name,
        )
        if not out_path:
            return

        # Export to a temp docx first, then convert
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_docx = tmp.name

        self._status_var.set("Building .docx for PDF conversion…")
        self.root.update_idletasks()

        ok = self._do_export_docx(tmp_docx, meta, open_after=False)
        if not ok:
            return

        self._status_var.set("Converting to PDF…")
        self.root.update_idletasks()

        log_lines: List[str] = []
        success = _export_pdf(tmp_docx, out_path, log=log_lines.append)

        try:
            Path(tmp_docx).unlink(missing_ok=True)
        except Exception:
            pass

        if success:
            self._status_var.set(f"PDF saved: {Path(out_path).name}")
            if messagebox.askyesno(APP_NAME, f"PDF saved:\n{out_path}\n\nOpen now?"):
                self._open_path(Path(out_path))
        else:
            detail = "\n".join(log_lines) or "No PDF converter found."
            messagebox.showerror(APP_NAME,
                f"PDF conversion failed.\n\n{detail}\n\n"
                "To enable PDF export, install one of:\n"
                "  pip install docx2pdf\n"
                "  — or —\n"
                "  Install LibreOffice (free) and ensure 'soffice' is on PATH.")
            self._status_var.set("PDF conversion failed — see error for options.")

    def _open_path(self, path: Path) -> None:
        try:
            if sys.platform == "win32":
                os.startfile(str(path))
            elif sys.platform == "darwin":
                subprocess.Popen(["open", str(path)])
            else:
                subprocess.Popen(["xdg-open", str(path)])
        except Exception:
            pass

    def _open_docs_dir(self) -> None:
        DOCS_DIR.mkdir(parents=True, exist_ok=True)
        self._open_path(DOCS_DIR)


# ── Entry point ───────────────────────────────────────────────────────────────
def _tk_runtime_help(err: Exception) -> str:
    lines = [
        f"{APP_NAME} cannot start because Tk/Tcl runtime is unavailable.",
        f"Python reported: {err}", "",
        "Remediation:",
    ]
    if sys.platform == "win32":
        lines.extend([
            "1. Repair/reinstall Python and include Tcl/Tk support.",
            "2. Verify this exists: <Python>\\tcl\\tcl8.6\\init.tcl",
            "3. Or run the packaged CITL executable build.",
        ])
    else:
        lines.extend([
            "1. Install tkinter (e.g. sudo apt install python3-tk).",
            "2. Restart the app.",
        ])
    return "\n".join(lines)


def main():
    try:
        root = tk.Tk()
    except tk.TclError as exc:
        msg = _tk_runtime_help(exc)
        log_path = _HERE / "citl_doc_composer_crash.log"
        try:
            log_path.write_text(msg + "\n\n" + traceback.format_exc(), encoding="utf-8")
        except Exception:
            pass
        print(msg, file=sys.stderr)
        sys.exit(2)

    root.withdraw()
    try:
        root.tk.call("tk", "scaling", 1.25)
    except Exception:
        pass
    root.deiconify()

    try:
        DocComposer(root)
        root.mainloop()
    except Exception as exc:
        log_path = _HERE / "citl_doc_composer_crash.log"
        log_path.write_text(f"{traceback.format_exc()}\n", encoding="utf-8")
        try:
            messagebox.showerror(APP_NAME, f"{exc}\nSee: {log_path}")
        except Exception:
            print(traceback.format_exc(), file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
