#!/usr/bin/env python3
"""
citl_doc_theme.py
CITL document print theme: colors, fonts, python-docx style application,
style presets, and font installation from repo/external font packs.
"""
from __future__ import annotations
import ctypes, os, re, shutil, sys
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

# ── Font pack location ────────────────────────────────────────────────────────
FONT_PACK = Path(r"C:\00 HENOSIS CODING PROJECTS\E READER REPO\fonts\reader-pack")

# Font family names as stored in the TTF files / used in python-docx
FONT_BODY     = "Berthold Baskerville"   # BertholdBaskerville.ttf
FONT_BODY_B   = "Berthold Baskerville"   # + bold=True
FONT_BODY_I   = "Berthold Baskerville"   # + italic=True
FONT_HEAD     = "Cheltenham"             # Cheltenham Book.ttf / Bold.ttf
FONT_CAPTION  = "Franklin Gothic Book"   # FranklinGothic Regular.ttf
FONT_MONO     = "Courier New"
FONT_FALLBACK = "Georgia"               # always present on Windows

# Map each logical name to its TTF file in the pack
FONT_FILES = {
    FONT_BODY:    [
        "BertholdBaskerville.ttf",
        "BertholdBaskerville-Bold.ttf",
        "BertholdBaskerville-Italic.ttf",
        "BertholdBaskerville-Book Italic.ttf",
    ],
    FONT_HEAD:    [
        "Cheltenham Book.ttf",
        "Cheltenham Bold.ttf",
        "Cheltenham BookItalic.ttf",
        "Cheltenham Italic.ttf",
    ],
    FONT_CAPTION: [
        "FranklinGothic Regular.ttf",
        "FranklinGothic Bold.ttf",
        "FranklinGothic Italic.ttf",
        "FranklinGothic Bold Italic.ttf",
    ],
}

_HERE = Path(__file__).resolve().parent
if getattr(sys, "frozen", False):
    _env_repo = os.environ.get("CITL_REPO", "").strip()
    if _env_repo and Path(_env_repo).is_dir():
        REPO = Path(_env_repo)
    else:
        REPO = Path(sys.executable).resolve().parent.parent.parent
else:
    REPO = _HERE.parent

# Repo-native font library for Doc Composer.
REPO_FONT_PACK = REPO / "factbook-assistant" / "fonts" / "doc_composer"

# External/legacy sources used for curated import fallback.
APOTHECARY_FONT_PACK = Path(
    r"M:\00 FONTS FONTS FONTS\QSL CARD FONTS\VINTAGE HEADER AND SUBHEADER\Apothecary Font Collection"
)

FONT_SOURCE_DIRS: Tuple[Path, ...] = (
    REPO_FONT_PACK,
    FONT_PACK,
    APOTHECARY_FONT_PACK,
)

DOC_STYLE_PRESETS: Dict[str, Dict[str, str]] = {
    "CITL Classic": {
        "body": "Berthold Baskerville",
        "heading": "Cheltenham",
        "caption": "Franklin Gothic Book",
        "fallback": "Georgia",
    },
    "State Grant Serif": {
        # Always-available Georgia — no font installation required.
        # Navy / maroon palette. Suitable for grant proposals and policy briefs.
        "body":     "Georgia",
        "heading":  "Georgia",
        "caption":  "Franklin Gothic Book",
        "fallback": "Georgia",
    },
    "Executive Sans": {
        "body": "Helvetica",
        "heading": "Avenir Next",
        "caption": "FF DIN",
        "fallback": "Arial",
    },
    "Humanist Professional": {
        "body": "Frutiger",
        "heading": "Avenir",
        "caption": "Trade Gothic",
        "fallback": "Arial",
    },
    "Editorial Modern": {
        "body": "Proxima Nova",
        "heading": "Futura",
        "caption": "Univers",
        "fallback": "Arial",
    },
    "Contemporary Clean": {
        "body": "Century Gothic",
        "heading": "Avenir Next",
        "caption": "Arial",
        "fallback": "Arial",
    },
    "Staff Walkthrough Blue": {
        "body": "Avenir Next",
        "heading": "Helvetica",
        "caption": "Arial",
        "fallback": "Arial",
    },
}
DOC_STYLE_NAMES: List[str] = list(DOC_STYLE_PRESETS.keys())
DEFAULT_DOC_STYLE = "Executive Sans"
_ACTIVE_DOC_STYLE = DEFAULT_DOC_STYLE

FONT_MATCH_RULES: Dict[str, List[Tuple[str, ...]]] = {
    "Berthold Baskerville": [("berthold", "baskerville"), ("baskerville",)],
    "Cheltenham": [("cheltenham",)],
    "Franklin Gothic Book": [("franklin", "gothic"), ("franklingothic",)],
    "Helvetica": [("helvetica",), ("lte5",)],
    "Avenir": [("avenir",)],
    "Avenir Next": [("avenirnext",), ("avenir", "next")],
    "FF DIN": [("dinpro",), ("ff", "din"), ("din", "next")],
    "Frutiger": [("frutiger",)],
    "Trade Gothic": [("trade", "gothic")],
    "Univers": [("univers",)],
    "Proxima Nova": [("proxima", "nova")],
    "Futura": [("futura",)],
    "Century Gothic": [("century", "gothic"), ("centurygothic",)],
    "Arial": [("arial",)],
    "Georgia": [("georgia",)],
}
FONT_FILE_EXTS = {".ttf", ".otf", ".ttc"}
MAX_FILES_PER_FAMILY = 40


def get_doc_style_names() -> List[str]:
    return list(DOC_STYLE_NAMES)


def _style_payload(style_name: Optional[str] = None) -> Dict[str, str]:
    key = (style_name or _ACTIVE_DOC_STYLE).strip()
    if key not in DOC_STYLE_PRESETS:
        key = DEFAULT_DOC_STYLE
    payload = dict(DOC_STYLE_PRESETS[key])
    payload["style"] = key
    return payload


def get_active_doc_style() -> str:
    return _ACTIVE_DOC_STYLE


def get_doc_style_fonts(style_name: Optional[str] = None) -> Dict[str, str]:
    return _style_payload(style_name)


def set_doc_style(style_name: Optional[str] = None) -> Dict[str, str]:
    global _ACTIVE_DOC_STYLE, FONT_BODY, FONT_BODY_B, FONT_BODY_I
    global FONT_HEAD, FONT_CAPTION, FONT_FALLBACK
    payload = _style_payload(style_name)
    _ACTIVE_DOC_STYLE = payload["style"]
    FONT_BODY = payload["body"]
    FONT_BODY_B = payload["body"]
    FONT_BODY_I = payload["body"]
    FONT_HEAD = payload["heading"]
    FONT_CAPTION = payload["caption"]
    FONT_FALLBACK = payload["fallback"]
    return payload


def get_required_families(style_name: Optional[str] = None, include_all_styles: bool = False) -> List[str]:
    if include_all_styles:
        families = set()
        for item in DOC_STYLE_PRESETS.values():
            families.add(item["body"])
            families.add(item["heading"])
            families.add(item["caption"])
        return sorted(families)
    payload = _style_payload(style_name)
    return [payload["body"], payload["heading"], payload["caption"]]


def get_missing_fonts(style_name: Optional[str] = None) -> List[str]:
    return [f for f in get_required_families(style_name) if not is_font_installed(f)]


# Initialize active style constants for import-time consumers.
set_doc_style(DEFAULT_DOC_STYLE)

# ── CITL print color palette ──────────────────────────────────────────────────
# Used as RGBColor(r, g, b) in python-docx calls
class _Palette:
    RED_ORANGE  = (0xCC, 0x33, 0x00)   # #CC3300  primary accent / rules
    SLATE_DARK  = (0x33, 0x4D, 0x6E)   # #334D6E  H1 / cover
    SLATE_MED   = (0x6B, 0x7F, 0x94)   # #6B7F94  H2 / sub-headers
    SLATE_LIGHT = (0xF0, 0xF3, 0xF6)   # #F0F3F6  callout background
    BODY_BLACK  = (0x1A, 0x1A, 0x1A)   # #1A1A1A  body text
    WHITE       = (0xFF, 0xFF, 0xFF)
    CAPTION     = (0x55, 0x55, 0x55)   # #555555  captions / footnotes
    RULE_HEX    = "CC3300"             # no-hash hex for XML attributes
    COVER_BG    = "334D6E"             # cover page header band

PAL = _Palette()

# ── Font detection (Windows registry, no pywin32 dep) ────────────────────────
def is_font_installed(family: str) -> bool:
    if sys.platform != "win32":
        return False
    import winreg
    keys = [
        (winreg.HKEY_LOCAL_MACHINE,
         r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
        (winreg.HKEY_CURRENT_USER,
         r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
    ]
    for hive, path in keys:
        try:
            with winreg.OpenKey(hive, path) as k:
                i = 0
                while True:
                    try:
                        name, _, _ = winreg.EnumValue(k, i)
                        if family.lower() in name.lower():
                            return True
                        i += 1
                    except OSError:
                        break
        except OSError:
            continue
    return False


def _iter_font_sources() -> Iterable[Path]:
    for src in FONT_SOURCE_DIRS:
        if src and src.exists():
            yield src


def _scan_source_font_files() -> List[Path]:
    files: List[Path] = []
    for src in _iter_font_sources():
        for p in src.rglob("*"):
            if p.is_file() and p.suffix.lower() in FONT_FILE_EXTS:
                files.append(p)
    return files


def _matches_family(file_name: str, family: str) -> bool:
    name = file_name.lower()
    rules = FONT_MATCH_RULES.get(family, [(family.lower(),)])
    for rule in rules:
        if all(token in name for token in rule):
            return True
    return False


def _family_candidates(family: str, all_files: List[Path]) -> List[Path]:
    picks = [p for p in all_files if _matches_family(p.name, family)]
    picks.sort(key=lambda p: (len(p.name), p.name.lower()))
    uniq: List[Path] = []
    seen = set()
    for p in picks:
        key = p.name.lower()
        if key in seen:
            continue
        seen.add(key)
        uniq.append(p)
    return uniq[:MAX_FILES_PER_FAMILY]


def install_citl_fonts(
    log=print,
    style_name: Optional[str] = None,
    include_all_styles: bool = True,
) -> dict:
    """
    Install Doc Composer fonts for current user (Windows, no admin).
    Sources are searched in repo pack first, then legacy/external locations.
    Returns {font_key: True/False} success map.
    """
    if sys.platform != "win32":
        log("[SKIP] Font install is Windows-only.")
        return {}

    source_roots = list(_iter_font_sources())
    if not source_roots:
        log("[WARN] No font source folders found.")
        return {}

    import winreg

    font_dir = Path(os.environ.get("LOCALAPPDATA", "")) / "Microsoft" / "Windows" / "Fonts"
    font_dir.mkdir(parents=True, exist_ok=True)

    all_files = _scan_source_font_files()
    families = get_required_families(style_name, include_all_styles=include_all_styles)
    results = {}
    installed_any = False

    for family in families:
        candidates = _family_candidates(family, all_files)
        if not candidates:
            key = f"{family}::(missing)"
            results[key] = False
            log(f"[MISS] {family}: no matching font files found in configured sources.")
            continue

        for src in candidates:
            dst = font_dir / src.name
            result_key = f"{family}::{src.name}"
            try:
                if not dst.exists():
                    shutil.copy2(src, dst)
                with winreg.OpenKey(
                    winreg.HKEY_CURRENT_USER,
                    r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts",
                    access=winreg.KEY_SET_VALUE,
                ) as k:
                    label_kind = "OpenType" if src.suffix.lower() == ".otf" else "TrueType"
                    label = f"{src.stem} ({label_kind})"
                    winreg.SetValueEx(k, label, 0, winreg.REG_SZ, str(dst))
                ctypes.windll.gdi32.AddFontResourceW(str(dst))
                results[result_key] = True
                installed_any = True
            except Exception as exc:
                results[result_key] = False
                log(f"[ERR] {result_key}: {exc}")

    if installed_any:
        try:
            ctypes.windll.user32.SendMessageW(0xFFFF, 0x001D, 0, 0)
        except Exception:
            pass
    return results


def resolve_font(preferred: str) -> str:
    """Return preferred font name if installed, else FONT_FALLBACK."""
    return preferred if is_font_installed(preferred) else FONT_FALLBACK


# ── python-docx style application ────────────────────────────────────────────
def apply_citl_styles(doc, style_name: Optional[str] = None) -> None:
    """
    Apply CITL print styles to a python-docx Document.
    Must be called before adding any content.
    """
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    payload = set_doc_style(style_name) if style_name else get_doc_style_fonts()
    body_font = resolve_font(payload["body"])
    heading_font = resolve_font(payload["heading"])
    caption_font = resolve_font(payload["caption"])
    is_staff_walkthrough = payload.get("style") == "Staff Walkthrough Blue"
    walkthrough_blue = (0x2E, 0x75, 0xB6)   # From "Staff Bot Creation Walkthrough" heading tone.
    walkthrough_gray = (0x55, 0x55, 0x55)

    # ---- Page layout --------------------------------------------------------
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin   = Cm(3.18)
    sec.right_margin  = Cm(2.54)

    # ---- Normal (body) ------------------------------------------------------
    normal = doc.styles["Normal"]
    nf = normal.font
    nf.name  = body_font
    nf.size  = Pt(11)
    nf.color.rgb = RGBColor(*PAL.BODY_BLACK)
    normal.paragraph_format.space_after  = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    # ---- Heading 1 ----------------------------------------------------------
    h1 = doc.styles["Heading 1"]
    h1f = h1.font
    h1f.name  = heading_font
    h1f.bold  = True
    h1f.size  = Pt(15.5 if is_staff_walkthrough else 18)
    h1f.color.rgb = RGBColor(*(walkthrough_blue if is_staff_walkthrough else PAL.SLATE_DARK))
    h1f.underline = False
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # ---- Heading 2 ----------------------------------------------------------
    h2 = doc.styles["Heading 2"]
    h2f = h2.font
    h2f.name  = heading_font
    h2f.bold  = False
    h2f.size  = Pt(12 if is_staff_walkthrough else 14)
    h2f.color.rgb = RGBColor(*(walkthrough_blue if is_staff_walkthrough else PAL.SLATE_MED))
    h2f.underline = False
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after  = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # ---- Heading 3 ----------------------------------------------------------
    h3 = doc.styles["Heading 3"]
    h3f = h3.font
    h3f.name   = caption_font
    h3f.bold   = True
    h3f.italic = False
    h3f.size   = Pt(11)
    h3f.color.rgb = RGBColor(*PAL.BODY_BLACK)
    h3f.underline = False
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after  = Pt(2)
    h3.paragraph_format.keep_with_next = True

    # ---- Caption style ------------------------------------------------------
    try:
        cap = doc.styles["Caption"]
    except KeyError:
        cap = doc.styles.add_style("Caption", 1)
    cap.font.name  = caption_font
    cap.font.size  = Pt(9.5 if is_staff_walkthrough else 9)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor(*(walkthrough_gray if is_staff_walkthrough else PAL.CAPTION))
    cap.paragraph_format.space_after = Pt(8)

    # ---- Header / Footer ----------------------------------------------------
    _build_header(doc, heading_font, caption_font)
    _build_footer(doc, caption_font)


def _build_header(doc, heading_font: str, caption_font: str) -> None:
    from docx.shared import Pt, RGBColor, Tab
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    header = doc.sections[0].header
    header.is_linked_to_previous = False
    # Clear default paragraph
    for p in header.paragraphs:
        p.clear()
    if not header.paragraphs:
        header.add_paragraph()

    para = header.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_citl = para.add_run("CITL  ")
    run_citl.font.name  = heading_font
    run_citl.font.bold  = True
    run_citl.font.size  = Pt(9)
    run_citl.font.color.rgb = RGBColor(*PAL.RED_ORANGE)

    run_title = para.add_run("Center for Information Technology and Learning")
    run_title.font.name  = caption_font
    run_title.font.size  = Pt(9)
    run_title.font.color.rgb = RGBColor(*PAL.CAPTION)

    # Red-orange bottom rule on the header paragraph
    _add_para_border(para, side="bottom", color=PAL.RULE_HEX, sz="6")


def _build_footer(doc, caption_font: str) -> None:
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    if not footer.paragraphs:
        footer.add_paragraph()

    para = footer.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_para_border(para, side="top", color=PAL.RULE_HEX, sz="4")

    # "Page N" field
    run = para.add_run()
    run.font.name  = caption_font
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(*PAL.CAPTION)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)

    run2 = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = para.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)

    para.add_run("  ·  CITL Documentation").font.size = Pt(9)


# ── DOCX building helpers ─────────────────────────────────────────────────────
def add_rule(doc, color_hex: str = PAL.RULE_HEX) -> None:
    """Add a thin colored horizontal rule paragraph."""
    from docx.shared import Pt
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    _add_para_border(para, side="bottom", color=color_hex, sz="6")


def add_h1_with_bar(doc, text: str) -> None:
    """H1 heading with a left red-orange bar."""
    para = doc.add_heading(text, level=1)
    _add_para_border(para, side="left",
                     color=PAL.RULE_HEX, sz="20", space="6")


def add_h2(doc, text: str) -> None:
    doc.add_heading(text, level=2)


def add_h3(doc, text: str) -> None:
    doc.add_heading(text, level=3)


def add_body(doc, text: str) -> None:
    """Add a Normal-style paragraph, auto-detecting numbered steps."""
    from docx.shared import Pt, Inches
    text = text.strip()
    if not text:
        return
    # Detect numbered step: "1." or "Step 1:"
    if re.match(r"^(\d+\.|\bStep\s+\d+:)", text):
        p = doc.add_paragraph(style="List Number")
        p.text = re.sub(r"^(\d+\.|\bStep\s+\d+:)\s*", "", text)
    elif re.match(r"^[-•]\s", text):
        p = doc.add_paragraph(style="List Bullet")
        p.text = text[2:]
    else:
        doc.add_paragraph(text)


def add_callout(doc, text: str, kind: str = "note") -> None:
    """
    Add a shaded callout box.
    kind: "note" | "warning" | "tip"
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    labels = {"note": "NOTE", "warning": "WARNING", "tip": "TIP"}
    label_colors = {
        "note":    PAL.SLATE_MED,
        "warning": (0xCC, 0x44, 0x00),
        "tip":     (0x22, 0x77, 0x44),
    }
    label = labels.get(kind, "NOTE")
    lcolor = label_colors.get(kind, PAL.SLATE_MED)

    # Use a 1-column table for the shaded box
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, "F0F3F6")
    _set_cell_border(cell, PAL.RULE_HEX if kind == "warning" else "6B7F94")

    # Label run
    p = cell.paragraphs[0]
    p.clear()
    label_run = p.add_run(label + "  ")
    label_run.font.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor(*lcolor)
    label_run.font.name = resolve_font(FONT_CAPTION)

    body_run = p.add_run(text)
    body_run.font.size = Pt(10)
    body_run.font.name = resolve_font(FONT_BODY)
    body_run.font.color.rgb = RGBColor(*PAL.BODY_BLACK)
    doc.add_paragraph()   # spacing after


def add_screenshot_placeholder(doc, caption: str = "", lines: int = 7) -> None:
    """
    Add a bordered screenshot placeholder box so editors can paste captures later.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, "FBFCFD")
    _set_cell_border(cell, "6B7F94")

    p = cell.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head = p.add_run("SCREENSHOT PLACEHOLDER")
    head.font.bold = True
    head.font.size = Pt(10)
    head.font.name = resolve_font(FONT_CAPTION)
    head.font.color.rgb = RGBColor(*PAL.SLATE_MED)

    cap = caption.strip() or "Paste UI capture here."
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_run = p2.add_run(cap)
    c_run.font.size = Pt(9)
    c_run.font.name = resolve_font(FONT_CAPTION)
    c_run.font.color.rgb = RGBColor(*PAL.CAPTION)

    for _ in range(max(4, int(lines))):
        cell.add_paragraph(" ")

    doc.add_paragraph()


def add_cover_page(doc, meta: dict) -> None:
    """
    Generate a styled CITL cover page.
    meta keys: title, subtitle, app_name, version, author, date, template_type
    """
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    style_name = (meta or {}).get("doc_style") if isinstance(meta, dict) else None
    payload = set_doc_style(style_name) if style_name else get_doc_style_fonts()
    heading_font = resolve_font(payload["heading"])
    body_font = resolve_font(payload["body"])
    caption_font = resolve_font(payload["caption"])

    def _centered(text, font, size, bold=False, color=PAL.BODY_BLACK, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        r.font.name  = font
        r.font.size  = Pt(size)
        r.font.bold  = bold
        r.font.color.rgb = RGBColor(*color)
        return p

    # Top spacer
    for _ in range(4):
        doc.add_paragraph()

    # CITL label
    _centered("CENTER FOR INFORMATION TECHNOLOGY AND LEARNING",
              caption_font, 9, color=PAL.RED_ORANGE, space_after=2)

    # Red rule
    add_rule(doc)
    doc.add_paragraph()

    # Document type
    _centered(meta.get("template_type", "Technical Document").upper(),
              caption_font, 10, color=PAL.SLATE_MED, space_after=18)

    # App name
    _centered(meta.get("app_name", "CITL Application"),
              heading_font, 28, bold=True, color=PAL.SLATE_DARK, space_after=6)

    # Title
    _centered(meta.get("title", ""),
              heading_font, 18, color=PAL.RED_ORANGE, space_after=4)

    # Subtitle
    if meta.get("subtitle"):
        _centered(meta["subtitle"], body_font, 12,
                  color=PAL.SLATE_MED, space_after=0)

    doc.add_paragraph()
    add_rule(doc)
    doc.add_paragraph()

    # Metadata block
    for label, key in [("Version", "version"), ("Author", "author"), ("Date", "date")]:
        val = meta.get(key, "")
        if val:
            _centered(f"{label}:  {val}", caption_font, 10, color=PAL.CAPTION)

    # Page break
    doc.add_page_break()


# ── Internal XML helpers ──────────────────────────────────────────────────────
def _add_para_border(para, side: str, color: str,
                     sz: str = "6", space: str = "1") -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr  = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"),   "single")
    bdr.set(qn("w:sz"),    sz)
    bdr.set(qn("w:space"), space)
    bdr.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(bdr)
    pPr.append(pBdr)


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_border(cell, color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBdr")
    for side in ("top", "bottom", "left", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"),   "single")
        bdr.set(qn("w:sz"),    "4")
        bdr.set(qn("w:space"), "0")
        bdr.set(qn("w:color"), color.lstrip("#"))
        tcBdr.append(bdr)
    tcPr.append(tcBdr)


# ── Grant / Proposal style helpers ────────────────────────────────────────────
# These are used by the Doc Composer when the "State Grant Serif" style is
# active and by the standalone _build_grant_doc.py builder script.

class _GrantPalette:
    """Navy / maroon / gold palette used for grant and policy documents."""
    NAVY        = (0x1C, 0x34, 0x54)   # deep navy — main headings
    MAROON      = (0x7B, 0x0C, 0x1A)   # RTC maroon — Part headings
    GOLD        = (0xB8, 0x8A, 0x00)   # gold accent
    BODY        = (0x1E, 0x1E, 0x1E)   # near-black body
    MUTED       = (0x4A, 0x4A, 0x4A)   # footer / captions
    WHITE       = (0xFF, 0xFF, 0xFF)
    NAVY_HEX    = "1C3454"
    MAROON_HEX  = "7B0C1A"
    INFO_FILL   = "EDF2F7"             # pale blue info-box background
    KW_FILL     = "F5F0E8"            # tan keyword-block background
    TBL_HEAD    = "1C3454"             # table header fill
    TBL_ALT     = "F0F4F8"            # alternating row

GRANT_PAL = _GrantPalette()
_GRANT_BODY_FONT = "Georgia"
_GRANT_HEAD_FONT = "Georgia"
_GRANT_MONO_FONT = "Courier New"
_GRANT_CAPT_FONT = "Franklin Gothic Book"


def _grant_font(family: str) -> str:
    """Return family if installed, else Georgia."""
    return family if is_font_installed(family) else "Georgia"


def add_grant_cover_page(
    doc,
    title: str,
    subtitle: str,
    institution: str = "Center for Instructional Technology and Learning (CITL)",
    college: str = "Renton Technical College",
    prepared_for: str = "",
    purpose: str = "",
) -> None:
    """
    Build a grant-style title page: large serif title, ruled divider,
    institution block, and a shaded "Prepared for / Purpose" box.
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    hf = _grant_font(_GRANT_HEAD_FONT)
    bf = _grant_font(_GRANT_BODY_FONT)

    def _cen(text, font, size, bold=False, italic=False,
             color=GRANT_PAL.NAVY, before=0, after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        r = p.add_run(text)
        r.font.name   = font
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(*color)
        return p

    # Spacer
    for _ in range(3):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    _cen(title,    hf, 22, bold=True, before=0, after=8)
    _cen(subtitle, hf, 13, italic=True, color=GRANT_PAL.MAROON, after=12)

    # Ruled divider
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(8)
    rule_p.paragraph_format.space_after  = Pt(8)
    _add_para_border(rule_p, side="bottom", color=GRANT_PAL.NAVY_HEX, sz="6")

    _cen(institution, hf, 11, bold=True, after=4)
    _cen(college,     hf, 11, bold=True, after=12)

    # "Prepared for / Purpose" info box
    if prepared_for or purpose:
        for label, body in [
            ("Prepared for: ", prepared_for),
            ("Document purpose: ", purpose),
        ]:
            if not body:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.space_after  = Pt(5)
            _set_para_shading(p, GRANT_PAL.INFO_FILL)
            rl = p.add_run(label)
            rl.font.name  = bf
            rl.font.size  = Pt(10)
            rl.font.bold  = True
            rl.font.color.rgb = RGBColor(*GRANT_PAL.NAVY)
            rb = p.add_run(body)
            rb.font.name  = bf
            rb.font.size  = Pt(10)
            rb.font.color.rgb = RGBColor(*GRANT_PAL.BODY)

    doc.add_page_break()


def _set_para_shading(para, fill_hex: str) -> None:
    """Apply a background shading fill to a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.lstrip("#"))
    pPr.append(shd)


def add_grant_part_heading(doc, text: str) -> None:
    """
    Part I / Part II heading — maroon all-caps with a bottom rule.
    """
    from docx.shared import Pt, RGBColor
    hf = _grant_font(_GRANT_HEAD_FONT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text.upper())
    r.font.name  = hf
    r.font.size  = Pt(13)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(*GRANT_PAL.MAROON)
    _add_para_border(p, side="bottom", color=GRANT_PAL.MAROON_HEX, sz="4")


def add_grant_app_banner(doc, text: str) -> None:
    """
    APP N: Title — navy filled banner with white text.
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    hf = _grant_font(_GRANT_HEAD_FONT)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0)
    _set_para_shading(p, GRANT_PAL.NAVY_HEX)
    r = p.add_run("  " + text + "  ")
    r.font.name  = hf
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(*GRANT_PAL.WHITE)


def add_grant_info_box(doc, label: str, body_text: str) -> None:
    """
    Shaded info / callout box with a bold navy label.
    Useful for "Prepared for:", "Note:", "Important:" blocks.
    """
    from docx.shared import Pt, RGBColor, Inches
    bf = _grant_font(_GRANT_BODY_FONT)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(8)
    _set_para_shading(p, GRANT_PAL.INFO_FILL)
    rl = p.add_run(label + "  ")
    rl.font.name  = bf
    rl.font.size  = Pt(10)
    rl.font.bold  = True
    rl.font.color.rgb = RGBColor(*GRANT_PAL.NAVY)
    rb = p.add_run(body_text)
    rb.font.name  = bf
    rb.font.size  = Pt(10)
    rb.font.color.rgb = RGBColor(*GRANT_PAL.BODY)


def add_grant_keyword_block(doc, keywords: str) -> None:
    """
    Tan-shaded keyword display block — bold navy label + Courier New keywords.
    Use for resume/job-board keyword sections.
    """
    from docx.shared import Pt, RGBColor, Inches
    bf = _grant_font(_GRANT_BODY_FONT)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    _set_para_shading(p, GRANT_PAL.KW_FILL)
    rl = p.add_run("Resume / Job-Board Keywords:  ")
    rl.font.name  = bf
    rl.font.size  = Pt(9.5)
    rl.font.bold  = True
    rl.font.color.rgb = RGBColor(*GRANT_PAL.NAVY)
    rb = p.add_run(keywords)
    rb.font.name  = _GRANT_MONO_FONT
    rb.font.size  = Pt(9)
    rb.font.color.rgb = RGBColor(*GRANT_PAL.MAROON)


def add_grant_skills_table(doc, rows: list) -> None:
    """
    Build a styled grant skills/alignment table.
    rows: list of lists of strings. First row is treated as header.
    """
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not rows:
        return
    n_cols = len(rows[0])
    table  = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    bf  = _grant_font(_GRANT_BODY_FONT)
    alt = (0xF0, 0xF4, 0xF8)

    for r_i, row_data in enumerate(rows):
        row = table.rows[r_i]
        for c_i, cell_text in enumerate(row_data):
            cell = row.cells[c_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Background fill
            fill = GRANT_PAL.TBL_HEAD if r_i == 0 else (GRANT_PAL.TBL_ALT if r_i % 2 == 1 else "FFFFFF")
            _set_cell_shading(cell, fill)
            para = cell.paragraphs[0]
            para.paragraph_format.left_indent = Inches(0.06)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            run = para.add_run(cell_text.strip())
            run.font.name  = bf
            run.font.size  = Pt(9.5)
            run.font.bold  = (r_i == 0)
            run.font.color.rgb = RGBColor(
                *(GRANT_PAL.WHITE if r_i == 0 else (0x1E, 0x1E, 0x1E))
            )
    doc.add_paragraph()
